"""
Document Processor Module for Domain Extraction Agent.

Enhanced document processing with support for multiple formats, intelligent chunking,
and metadata extraction specifically for technical documentation and certification guides.
"""

import asyncio
import re
import json
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from datetime import datetime
import hashlib

from loguru import logger
import pypdf  # Modern version of PyPDF2
import markdown
from bs4 import BeautifulSoup
import docx
import ebooklib
from ebooklib import epub

from .models import (
    Document,
    DocumentChunk,
    DocumentType,
    ConceptType,
    DomainCategory
)
from ....core.utils import clean_text


class DocumentProcessor:
    """Enhanced document processor for domain extraction."""
    
    def __init__(self):
        self.default_chunk_size = 500
        self.default_overlap = 50
        self.min_chunk_size = 100
        self.sentence_splitter = re.compile(r'(?<=[.!?])\s+')
        self.section_patterns = {
            'markdown': re.compile(r'^#{1,6}\s+(.+)$', re.MULTILINE),
            'general': re.compile(r'^(?:Chapter|Section|Part)\s*\d*\s*[:\-]?\s*(.+)$', re.MULTILINE | re.IGNORECASE)
        }
        
    async def process_documents(
        self,
        file_paths: List[str],
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None
    ) -> Tuple[List[Document], List[DocumentChunk]]:
        """Process multiple documents and return documents and chunks."""
        try:
            logger.info(f"Processing {len(file_paths)} documents")
            
            documents = []
            all_chunks = []
            
            for file_path in file_paths:
                doc, chunks = await self.process_single_document(
                    file_path, chunk_size, chunk_overlap
                )
                if doc:
                    documents.append(doc)
                    all_chunks.extend(chunks)
                    
            logger.info(f"Processed {len(documents)} documents into {len(all_chunks)} chunks")
            return documents, all_chunks
            
        except Exception as e:
            logger.error(f"Error processing documents: {str(e)}")
            raise
            
    async def process_single_document(
        self,
        file_path: str,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None
    ) -> Tuple[Optional[Document], List[DocumentChunk]]:
        """Process a single document file."""
        try:
            path = Path(file_path)
            if not path.exists():
                logger.warning(f"File not found: {file_path}")
                return None, []
                
            # Determine document type
            doc_type = self._get_document_type(path)
            
            # Extract content based on type
            if doc_type == DocumentType.PDF:
                content, metadata = await self._process_pdf(path)
            elif doc_type == DocumentType.MARKDOWN:
                content, metadata = await self._process_markdown(path)
            elif doc_type == DocumentType.DOCX:
                content, metadata = await self._process_docx(path)
            elif doc_type == DocumentType.HTML:
                content, metadata = await self._process_html(path)
            elif doc_type == DocumentType.EPUB:
                content, metadata = await self._process_epub(path)
            else:
                content, metadata = await self._process_text(path)
                
            if not content:
                logger.warning(f"No content extracted from: {file_path}")
                return None, []
                
            # Extract sections
            sections = self._extract_sections(content, doc_type)
            
            # Create document
            document = Document(
                source_path=str(path),
                document_type=doc_type,
                title=metadata.get('title', path.stem),
                content=content,
                metadata=metadata,
                sections=sections,
                word_count=len(content.split()),
                page_count=metadata.get('page_count')
            )
            
            # Create chunks
            chunks = await self._create_chunks(
                document,
                chunk_size or self.default_chunk_size,
                chunk_overlap or self.default_overlap
            )
            
            return document, chunks
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            return None, []
            
    def _get_document_type(self, path: Path) -> DocumentType:
        """Determine document type from file extension."""
        ext = path.suffix.lower()
        type_map = {
            '.pdf': DocumentType.PDF,
            '.md': DocumentType.MARKDOWN,
            '.markdown': DocumentType.MARKDOWN,
            '.txt': DocumentType.TEXT,
            '.html': DocumentType.HTML,
            '.htm': DocumentType.HTML,
            '.docx': DocumentType.DOCX,
            '.epub': DocumentType.EPUB
        }
        return type_map.get(ext, DocumentType.TEXT)
        
    async def _process_pdf(self, path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process PDF document."""
        try:
            content_parts = []
            metadata = {
                'filename': path.name,
                'file_type': 'pdf',
                'processed_date': datetime.utcnow().isoformat()
            }
            
            with open(path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                metadata['page_count'] = len(pdf_reader.pages)
                
                # Extract metadata
                if pdf_reader.metadata:
                    metadata['pdf_metadata'] = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', '')
                    }
                    
                # Extract text from each page
                for i, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text:
                        content_parts.append(f"[Page {i+1}]\n{text}")
                        
            content = '\n\n'.join(content_parts)
            
            # Extract title from content if not in metadata
            if not metadata.get('pdf_metadata', {}).get('title'):
                metadata['title'] = self._extract_title(content)
                
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            raise
            
    async def _process_markdown(self, path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process Markdown document."""
        try:
            content = path.read_text(encoding='utf-8')
            
            # Parse markdown
            md = markdown.Markdown(extensions=['meta', 'tables', 'fenced_code'])
            html = md.convert(content)
            
            # Extract metadata
            metadata = {
                'filename': path.name,
                'file_type': 'markdown',
                'processed_date': datetime.utcnow().isoformat()
            }
            
            if hasattr(md, 'Meta'):
                metadata['markdown_metadata'] = md.Meta
                
            # Extract plain text from HTML
            soup = BeautifulSoup(html, 'html.parser')
            plain_text = soup.get_text()
            
            # Keep original markdown for better section extraction
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error processing Markdown: {str(e)}")
            raise
            
    async def _process_docx(self, path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process DOCX document."""
        try:
            doc = docx.Document(path)
            
            content_parts = []
            metadata = {
                'filename': path.name,
                'file_type': 'docx',
                'processed_date': datetime.utcnow().isoformat()
            }
            
            # Extract core properties
            core_props = doc.core_properties
            metadata['docx_metadata'] = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': core_props.created.isoformat() if core_props.created else None
            }
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)
                    
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(' | '.join(row_text))
                if table_text:
                    content_parts.append('\n'.join(table_text))
                    
            content = '\n\n'.join(content_parts)
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise
            
    async def _process_html(self, path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process HTML document."""
        try:
            content = path.read_text(encoding='utf-8')
            
            soup = BeautifulSoup(content, 'html.parser')
            
            # Extract metadata
            metadata = {
                'filename': path.name,
                'file_type': 'html',
                'processed_date': datetime.utcnow().isoformat()
            }
            
            # Extract title
            title_tag = soup.find('title')
            if title_tag:
                metadata['title'] = title_tag.text.strip()
                
            # Extract meta tags
            meta_data = {}
            for meta in soup.find_all('meta'):
                if meta.get('name'):
                    meta_data[meta.get('name')] = meta.get('content', '')
            if meta_data:
                metadata['html_metadata'] = meta_data
                
            # Extract text content
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
                
            text = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            content = '\n'.join(chunk for chunk in chunks if chunk)
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error processing HTML: {str(e)}")
            raise
            
    async def _process_epub(self, path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process EPUB document."""
        try:
            book = epub.read_epub(path)
            
            content_parts = []
            metadata = {
                'filename': path.name,
                'file_type': 'epub',
                'processed_date': datetime.utcnow().isoformat()
            }
            
            # Extract metadata
            metadata['epub_metadata'] = {
                'title': book.get_metadata('DC', 'title')[0][0] if book.get_metadata('DC', 'title') else '',
                'creator': book.get_metadata('DC', 'creator')[0][0] if book.get_metadata('DC', 'creator') else '',
                'language': book.get_metadata('DC', 'language')[0][0] if book.get_metadata('DC', 'language') else ''
            }
            
            # Extract chapters
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    text = soup.get_text()
                    if text.strip():
                        content_parts.append(text)
                        
            content = '\n\n'.join(content_parts)
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error processing EPUB: {str(e)}")
            raise
            
    async def _process_text(self, path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process plain text document."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'windows-1252']
            content = None
            
            for encoding in encodings:
                try:
                    content = path.read_text(encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
                    
            if content is None:
                raise ValueError(f"Could not decode file with any encoding: {encodings}")
                
            metadata = {
                'filename': path.name,
                'file_type': 'text',
                'processed_date': datetime.utcnow().isoformat(),
                'title': self._extract_title(content)
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error processing text file: {str(e)}")
            raise
            
    def _extract_sections(self, content: str, doc_type: DocumentType) -> List[Dict[str, str]]:
        """Extract sections from content."""
        sections = []
        
        # Choose pattern based on document type
        if doc_type == DocumentType.MARKDOWN:
            pattern = self.section_patterns['markdown']
        else:
            pattern = self.section_patterns['general']
            
        # Find all section headers
        matches = list(pattern.finditer(content))
        
        if not matches:
            # No sections found, treat entire content as one section
            return [{
                'title': 'Main Content',
                'content': content[:1000],  # Preview only
                'start_pos': 0,
                'end_pos': len(content)
            }]
            
        # Extract sections
        for i, match in enumerate(matches):
            title = match.group(1).strip()
            start_pos = match.end()
            
            # Find end position (start of next section or end of content)
            if i < len(matches) - 1:
                end_pos = matches[i + 1].start()
            else:
                end_pos = len(content)
                
            section_content = content[start_pos:end_pos].strip()
            
            sections.append({
                'title': title,
                'content': section_content[:1000],  # Preview only
                'start_pos': start_pos,
                'end_pos': end_pos
            })
            
        return sections
        
    def _extract_title(self, content: str) -> str:
        """Extract title from content."""
        lines = content.split('\n')
        
        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            if line and len(line) < 200:  # Reasonable title length
                # Skip lines that look like metadata
                if not any(char in line for char in ['=', ':', '|', '#']):
                    return line
                    
        return "Untitled Document"
        
    async def _create_chunks(
        self,
        document: Document,
        chunk_size: int,
        chunk_overlap: int
    ) -> List[DocumentChunk]:
        """Create chunks from document with intelligent splitting."""
        try:
            chunks = []
            
            # Split content into sentences
            sentences = self._split_into_sentences(document.content)
            
            if not sentences:
                return chunks
                
            current_chunk = []
            current_size = 0
            start_char = 0
            current_start = 0
            
            for i, sentence in enumerate(sentences):
                sentence_words = sentence.split()
                sentence_size = len(sentence_words)
                
                # Check if adding this sentence exceeds chunk size
                if current_size + sentence_size > chunk_size and current_chunk:
                    # Create chunk
                    chunk_content = ' '.join(current_chunk)
                    chunk = DocumentChunk(
                        document_id=document.id,
                        content=chunk_content,
                        chunk_index=len(chunks),
                        total_chunks=0,  # Will be updated later
                        start_char=current_start,
                        end_char=start_char,
                        metadata={
                            'source_file': document.source_path,
                            'document_title': document.title
                        }
                    )
                    chunks.append(chunk)
                    
                    # Prepare overlap for next chunk
                    if chunk_overlap > 0:
                        # Calculate overlap sentences
                        overlap_size = 0
                        overlap_sentences = []
                        
                        for j in range(len(current_chunk) - 1, -1, -1):
                            sent_size = len(current_chunk[j].split())
                            if overlap_size + sent_size <= chunk_overlap:
                                overlap_sentences.insert(0, current_chunk[j])
                                overlap_size += sent_size
                            else:
                                break
                                
                        current_chunk = overlap_sentences
                        current_size = overlap_size
                        current_start = start_char - len(' '.join(overlap_sentences))
                    else:
                        current_chunk = []
                        current_size = 0
                        current_start = start_char
                        
                # Add sentence to current chunk
                current_chunk.append(sentence)
                current_size += sentence_size
                start_char += len(sentence) + 1  # +1 for space
                
            # Don't forget the last chunk
            if current_chunk:
                chunk_content = ' '.join(current_chunk)
                chunk = DocumentChunk(
                    document_id=document.id,
                    content=chunk_content,
                    chunk_index=len(chunks),
                    total_chunks=0,
                    start_char=current_start,
                    end_char=len(document.content),
                    metadata={
                        'source_file': document.source_path,
                        'document_title': document.title
                    }
                )
                chunks.append(chunk)
                
            # Update total chunks count
            total = len(chunks)
            for chunk in chunks:
                chunk.total_chunks = total
                
            # Extract potential concepts from each chunk
            for chunk in chunks:
                chunk.concepts = self._extract_chunk_concepts(chunk.content)
                
            return chunks
            
        except Exception as e:
            logger.error(f"Error creating chunks: {str(e)}")
            raise
            
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences with improved handling."""
        # Handle common abbreviations
        text = re.sub(r'\b(Dr|Mr|Mrs|Ms|Prof|Sr|Jr)\.\s*', r'\1<PERIOD> ', text)
        text = re.sub(r'\b(Inc|Ltd|Corp|Co)\.\s*', r'\1<PERIOD> ', text)
        text = re.sub(r'\b(Jan|Feb|Mar|Apr|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)\.\s*', r'\1<PERIOD> ', text)
        
        # Handle numbers with periods
        text = re.sub(r'(\d+)\.(\d+)', r'\1<DECIMAL>\2', text)
        text = re.sub(r'(\d+)\.\s+(\d+)', r'\1<PERIOD> \2', text)
        
        # Split sentences
        sentences = self.sentence_splitter.split(text)
        
        # Restore periods
        sentences = [s.replace('<PERIOD>', '.').replace('<DECIMAL>', '.') for s in sentences]
        
        # Filter out empty sentences
        sentences = [s.strip() for s in sentences if s.strip()]
        
        return sentences
        
    def _extract_chunk_concepts(self, chunk_text: str) -> List[str]:
        """Extract potential concepts from chunk text."""
        concepts = []
        
        # Look for capitalized terms (potential service/feature names)
        cap_pattern = re.compile(r'\b[A-Z][a-zA-Z]*(?:\s+[A-Z][a-zA-Z]*)*\b')
        capitalized = cap_pattern.findall(chunk_text)
        
        # Look for technical terms in quotes
        quote_pattern = re.compile(r'["\']([^"\']+)["\']')
        quoted = quote_pattern.findall(chunk_text)
        
        # Look for acronyms
        acronym_pattern = re.compile(r'\b[A-Z]{2,}\b')
        acronyms = acronym_pattern.findall(chunk_text)
        
        # Combine and deduplicate
        all_concepts = set(capitalized + quoted + acronyms)
        
        # Filter out common words
        common_words = {'The', 'This', 'That', 'These', 'Those', 'A', 'An', 'And', 'Or', 'But', 'In', 'On', 'At', 'To', 'For'}
        concepts = [c for c in all_concepts if c not in common_words and len(c) > 2]
        
        return concepts[:10]  # Limit to top 10 per chunk
        
    async def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF file - compatibility method."""
        try:
            path = Path(pdf_path)
            content, _ = await self._process_pdf(path)
            return content
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            return ""
            
    async def chunk_text(
        self,
        text: str,
        chunk_size: int = None,
        overlap: int = None
    ) -> List[str]:
        """Chunk text into overlapping segments - compatibility method."""
        chunk_size = chunk_size or self.default_chunk_size
        overlap = overlap or self.default_overlap
        
        # Create temporary document
        temp_doc = Document(
            source_path="temp",
            document_type=DocumentType.TEXT,
            title="Temporary",
            content=text
        )
        
        # Create chunks
        chunks = await self._create_chunks(temp_doc, chunk_size, overlap)
        
        # Return just the text content
        return [chunk.content for chunk in chunks]
